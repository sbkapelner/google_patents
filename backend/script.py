import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches
from hyperlink import add_hyperlink


class requestFunctions:
    def get_html(self, result_link: str):
        my_headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 11_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1.2 Safari/605.1.15",
            "Accept-Encoding": "gzip, deflate",
            "Accept": "*/*",
            "Connection": "keep-alive",
        }
        response = requests.get(result_link, headers=my_headers)
        status_code = response.status_code

        return response, status_code

    def get_other(self, response: object) -> str:
        soup = BeautifulSoup(response.text, "html.parser")
        titleandid = soup.find_all("title")[0].text
        id = re.split(" - | \n", titleandid)[0]
        title = re.split(" - | \n", titleandid)[1]

        inventor_tag = soup.find("dd", itemprop="inventor")
        inventor = "Inventor: " + inventor_tag.text

        assignee_tag = soup.find("span", itemprop="assigneeSearch")
        if assignee_tag.text != inventor_tag.text:
            assignee = "Assignee: " + assignee_tag.text
        else:
            assignee = "Assignee: None"

        priority_date = soup.find(itemprop="priorityDate").text

        return (
            f"{id}\n",
            f"{title.strip()}\n",
            f"{inventor}\n",
            f"{assignee}\n",
            f"{priority_date}\n",
        )

    def get_status(self, response: object) -> str:
        soup = BeautifulSoup(response.text, "html.parser")
        status_tag = soup.find_all("span", itemprop="status")
        exp_tag = soup.find_all("time", itemprop="expiration")

        if not status_tag:
            status = 0
        else:
            status = status_tag[0].text
            if (
                "Expired" in status
            ):  # reset the status to Expired instead of "Expired - Fee Related"
                status = "Expired"

        if not exp_tag:
            exp = "DATE MISSING"
        else:
            exp = exp_tag[0].text

        match status:
            case "Expired":
                return f"{status}\n"
            case "Active":
                return f"{status} (Exp. {exp})\n"
            case 0:
                return "No status or exp date"
            case "Abandoned":
                return f"{status}\n"
            case "Pending":
                return f"{status}\n"
            case "Withdrawn":
                return f"{status}\n"

    def url_from_patno(self, no: str) -> str:
        url = f"https://patents.google.com/patent/{no}/en?oq={no}"

        return url


# file = updateDocument("dates.docx", result_link=)
# file.create_document()
# file.add_table()
# file.new_row(0)
class updateDocument(requestFunctions):
    def __init__(
        self,
        docxfilename: str,
        parameters={
            "patent_no": "on",
            "title": "on",
            "inventor": "on",
            "assignee": "on",
            "status": "on",
            "priority_date": "off",
        },
    ):
        self.docxfilename = docxfilename
        self.parameters = parameters

    def get_google_data(self, response):
        google_data = {
            "patent_no": self.get_other(response)[0],
            "title": self.get_other(response)[1],
            "inventor": self.get_other(response)[2],
            "assignee": self.get_other(response)[3],
            "status": self.get_status(response),
            "priority_date": self.get_other(response)[4],
        }
        return google_data

    def create_document(self):
        document = Document()
        document.save(self.docxfilename)

    def add_table(self):
        document = Document(self.docxfilename)
        document.add_paragraph()
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        document.save(self.docxfilename)

    def new_row(self, table_no: int, pat_no: str):
        result_link = self.url_from_patno(pat_no)
        response, status_code = self.get_html(result_link)

        document = Document(self.docxfilename)
        table = document.tables[table_no]
        table.add_row()
        cell = table._cells[len(table._cells) - 1]
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Inches(0.3)
        filtered_keys = [k for (k, v) in self.parameters.items() if v != "off"]

        if status_code == 200:
            for key in filtered_keys:
                if key == "patent_no":
                    add_hyperlink(
                        p, self.get_google_data(response)["patent_no"], result_link
                    )
                else:
                    p.add_run(self.get_google_data(response)[key])

        else:
            p.add_run(f"Could not find {pat_no}")

        document.save(self.docxfilename)
