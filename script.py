import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches
from hyperlink import add_hyperlink


class requestFunctions:
    def get_html(self, result_link: str):
        my_headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 11_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1.2 Safari/605.1.15",
            "Accept-Encoding": "gzip, deflate",
            "Accept": "*/*",
            "Connection": "keep-alive",
        }
        html = requests.get(result_link, headers=my_headers).text
        soup = BeautifulSoup(html, "html.parser")

        return soup

    def get_other(self, result_link: str) -> str:
        soup = self.get_html(result_link)
        titleandid = soup.find_all("title")[0].text
        id = re.split(" - | \n", titleandid)[0]
        title = re.split(" - | \n", titleandid)[1]

        inventor_tag = soup.find("dd", itemprop="inventor")
        inventor = "Inventor: " + inventor_tag.text

        assignee_tag = soup.find("span", itemprop="assigneeSearch")
        if assignee_tag.text != inventor_tag.text:
            assignee = "Assignee: " + assignee_tag.text
        else:
            assignee = "Assignee: None"

        priority_date = soup.find(itemprop="priorityDate").text

        return (
            f"{id}\n",
            f"{title.strip()}\n",
            f"{inventor}\n",
            f"{assignee}\n",
            f"{priority_date}\n",
        )

    def get_status(self, result_link: str):
        soup = self.get_html(result_link)
        status_tag = soup.find_all("span", itemprop="status")
        exp_tag = soup.find_all("time", itemprop="expiration")

        if not status_tag:
            status = 0
        else:
            status = status_tag[0].text
            if (
                "Expired" in status
            ):  # reset the status to Expired instead of "Expired - Fee Related"
                status = "Expired"

        if not exp_tag:
            exp = "DATE MISSING"
        else:
            exp = exp_tag[0].text

        match status:
            case "Expired":
                return f"{status}\n"
            case "Active":
                return f"{status} (Exp. {exp})\n"
            case 0:
                return "No status or exp date"
            case "Abandoned":
                return f"{status}\n"
            case "Pending":
                return f"{status}\n"
            case "Withdrawn":
                return f"{status}\n"

    def url_from_patno(self, no: str) -> str:
        url = f"https://patents.google.com/patent/{no}/en?oq={no}"

        return url


# file = updateDocument("dates.docx", result_link=)
# file.create_document()
# file.add_table()
# file.new_row(0)
class updateDocument(requestFunctions):
    def __init__(
        self,
        docxfilename: str,
        pat_no: str,
        parameters={
            "patent_no": "on",
            "title": "on",
            "inventor": "on",
            "assignee": "on",
            "status": "on",
            "priority_date": "off",
        },
    ):
        self.docxfilename = docxfilename
        self.result_link = self.url_from_patno(pat_no)
        self.parameters = parameters
        self.google_data = {
            "id": self.get_other(self.result_link)[0],
            "title": self.get_other(self.result_link)[1],
            "inventor": self.get_other(self.result_link)[2],
            "assignee": self.get_other(self.result_link)[3],
            "status": self.get_status(self.result_link),
            "prioritydate": self.get_other(self.result_link)[4],
        }

    def create_document(self):
        document = Document()
        document.save(self.docxfilename)

    def add_table(self):
        document = Document(self.docxfilename)
        document.add_paragraph()
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        document.save(self.docxfilename)

    def new_row(self, table_no: int):
        document = Document(self.docxfilename)
        table = document.tables[table_no]
        table.add_row()
        cell = table._cells[len(table._cells) - 1]
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Inches(0.3)

        for key_pair in list(zip(self.parameters, self.google_data)):
            if self.parameters[key_pair[0]] == "on" and key_pair[0] == "patent_no":
                add_hyperlink(p, self.get_other(self.result_link)[0], self.result_link)
            if self.parameters[key_pair[0]] == "on" and key_pair[0] != "patent_no":
                p.add_run(self.google_data[key_pair[1]])

        document.save(self.docxfilename)
